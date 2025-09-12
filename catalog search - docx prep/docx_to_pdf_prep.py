# -*- coding: utf-8 -*-
"""
DOCX/PDF to PDF Preprocessor for Catalog Search Pipeline

This script scans a folder on NAS for DOCX and PDF files,
converts DOCX files to PDF format, and saves all PDFs
to the output folder for processing by catalog search stage 1.

DOCX files are converted to PDF using docx2pdf which preserves
all original formatting, styles, images, tables, and layout.
PDF files are copied directly to the output folder without modification.
"""

import sys
import os
import platform
from smb.SMBConnection import SMBConnection
from smb import smb_structs
import io
import socket
from datetime import datetime
import tempfile
import shutil
try:
    # Try docx2pdf first (Windows/Mac with MS Word)
    from docx2pdf import convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

# LibreOffice via subprocess (best free cross-platform solution)
import subprocess
HAS_SUBPROCESS = True

try:
    # Optional: Aspose.Words for Python (commercial, best pure Python solution)
    import aspose.words as aw
    HAS_ASPOSE = True
except ImportError:
    HAS_ASPOSE = False
    
try:
    # Fallback to pypandoc (poor formatting preservation, last resort only)
    import pypandoc
    HAS_PYPANDOC = True
except ImportError:
    HAS_PYPANDOC = False

try:
    # Pure Python fallback - docx + fpdf2 for basic conversion
    from docx import Document
    from fpdf import FPDF
    HAS_DOCX_FPDF = True
except ImportError:
    HAS_DOCX_FPDF = False

try:
    # Alternative pure Python - python-docx + pdfkit (requires wkhtmltopdf)
    import pdfkit
    HAS_PDFKIT = True
except ImportError:
    HAS_PDFKIT = False

# ==============================================================================
# --- Configuration ---
# ==============================================================================

# --- NAS Configuration ---
# Network attached storage connection parameters
# NOTE: Update these values to match your actual NAS configuration
# These should match the configuration in catalog search/stage1_extract_csv.py
NAS_PARAMS = {
    "ip": "your_nas_ip",
    "share": "your_share_name",
    "user": "your_nas_user",
    "password": "your_nas_password",
    "port": 445
}

# Input path on NAS where DOCX/PDF files are located (relative to share)
# This folder will be scanned recursively for all .docx, .doc, and .pdf files
NAS_INPUT_FOLDER_PATH = "path/to/docx_pdf/input/folder"  # e.g., "iris/documents/input"

# Output path on NAS where PDFs will be saved (relative to share)
# All PDFs (converted and copied) will be saved here with original filenames
NAS_OUTPUT_FOLDER_PATH = "path/to/pdf/output/folder"  # e.g., "iris/documents/output"

# --- pysmb Configuration ---
smb_structs.SUPPORT_SMB2 = True
smb_structs.MAX_PAYLOAD_SIZE = 65536
CLIENT_HOSTNAME = socket.gethostname()

# ==============================================================================
# --- Helper Functions ---
# ==============================================================================

def create_nas_connection():
    """Creates and returns an authenticated SMBConnection object."""
    try:
        conn = SMBConnection(
            NAS_PARAMS["user"],
            NAS_PARAMS["password"],
            CLIENT_HOSTNAME,
            NAS_PARAMS["ip"],
            use_ntlm_v2=True,
            is_direct_tcp=(NAS_PARAMS["port"] == 445)
        )
        connected = conn.connect(NAS_PARAMS["ip"], NAS_PARAMS["port"], timeout=60)
        if not connected:
            print("[ERROR] Failed to connect to NAS.")
            return None
        print(f"Successfully connected to NAS: {NAS_PARAMS['ip']}:{NAS_PARAMS['port']} on share '{NAS_PARAMS['share']}'")
        return conn
    except Exception as e:
        print(f"[ERROR] Exception creating NAS connection: {e}")
        return None

def ensure_nas_dir_exists(conn, share_name, dir_path):
    """Ensures a directory exists on the NAS, creating it if necessary."""
    if not conn:
        print("[ERROR] Cannot ensure NAS directory: No connection.")
        return False
    
    path_parts = dir_path.strip('/').split('/')
    current_path = ''
    try:
        for part in path_parts:
            if not part: continue
            current_path = os.path.join(current_path, part).replace('\\', '/')
            try:
                conn.listPath(share_name, current_path)
            except Exception:
                print(f"Creating directory on NAS: {current_path}")
                conn.createDirectory(share_name, current_path)
        return True
    except Exception as e:
        print(f"[ERROR] Failed to ensure/create NAS directory '{dir_path}': {e}")
        return False

def list_docx_pdf_files_from_nas(conn, share_name, folder_path):
    """Recursively list all DOCX and PDF files in the specified NAS folder."""
    files_list = []
    
    def walk_directory(current_path):
        try:
            items = conn.listPath(share_name, current_path)
            for item in items:
                if item.filename == '.' or item.filename == '..':
                    continue
                
                # Skip system/temporary files
                if item.filename.startswith('~$') or item.filename.startswith('.'):
                    print(f"Skipping system/temporary file: {item.filename}")
                    continue
                
                full_path = os.path.join(current_path, item.filename).replace('\\', '/')
                
                if item.isDirectory:
                    # Recursively walk subdirectories
                    walk_directory(full_path)
                else:
                    # Check if it's a DOCX, DOC, or PDF file
                    lower_filename = item.filename.lower()
                    if lower_filename.endswith(('.docx', '.doc', '.pdf')):
                        files_list.append({
                            'path': full_path,
                            'filename': item.filename,
                            'size': item.file_size,
                            'is_pdf': lower_filename.endswith('.pdf')
                        })
        except Exception as e:
            print(f"[WARNING] Error listing directory '{current_path}': {e}")
    
    try:
        # Start walking from the root folder
        walk_directory(folder_path)
        return files_list
    except Exception as e:
        print(f"[ERROR] Failed to list files from NAS: {e}")
        return []

def read_file_from_nas(conn, share_name, file_path):
    """Read a file from NAS and return its content as bytes."""
    try:
        file_obj = io.BytesIO()
        file_attributes, filesize = conn.retrieveFile(share_name, file_path, file_obj)
        file_obj.seek(0)
        return file_obj.read()
    except Exception as e:
        print(f"[ERROR] Failed to read file from NAS '{file_path}': {e}")
        return None

def save_pdf_to_nas(conn, share_name, file_path, pdf_content):
    """Save a PDF to NAS."""
    try:
        # Ensure directory exists
        dir_path = os.path.dirname(file_path).replace('\\', '/')
        if dir_path and not ensure_nas_dir_exists(conn, share_name, dir_path):
            print(f"[ERROR] Failed to ensure output directory exists: {dir_path}")
            return False
        
        # Write PDF to NAS
        file_obj = io.BytesIO(pdf_content)
        bytes_written = conn.storeFile(share_name, file_path, file_obj)
        print(f"Successfully wrote PDF ({bytes_written} bytes) to: {file_path}")
        return True
    except Exception as e:
        print(f"[ERROR] Failed to write PDF to NAS '{file_path}': {e}")
        return False

def copy_pdf_to_output(conn, share_name, source_path, dest_path):
    """Copy a PDF file from source to destination on NAS."""
    try:
        # Read the PDF content
        pdf_content = read_file_from_nas(conn, share_name, source_path)
        if pdf_content is None:
            return False
        
        # Save to destination
        return save_pdf_to_nas(conn, share_name, dest_path, pdf_content)
    except Exception as e:
        print(f"[ERROR] Failed to copy PDF from '{source_path}' to '{dest_path}': {e}")
        return False

def convert_docx_to_pdf(docx_content, filename):
    """
    Convert DOCX content to PDF format preserving all original formatting.
    
    Tries multiple methods in order of quality:
    1. docx2pdf (excellent quality, requires MS Word on Windows/Mac)
    2. Aspose.Words (excellent quality, commercial, pure Python)
    3. LibreOffice (excellent quality, free, cross-platform)
    4. pypandoc (poor formatting, last resort only)
    """
    
    # Create a temporary directory for file operations
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Write DOCX content to temporary file
        temp_docx_path = os.path.join(temp_dir, filename)
        temp_pdf_path = os.path.join(temp_dir, os.path.splitext(filename)[0] + '.pdf')
        
        with open(temp_docx_path, 'wb') as f:
            f.write(docx_content)
        
        # Method 1: Try docx2pdf (excellent quality, MS Word required)
        if HAS_DOCX2PDF:
            try:
                print(f"  Using docx2pdf for conversion (excellent quality)...")
                convert(temp_docx_path, temp_pdf_path)
                
                # Read the converted PDF
                with open(temp_pdf_path, 'rb') as f:
                    pdf_content = f.read()
                return pdf_content
            except Exception as e:
                print(f"  docx2pdf failed: {e}")
                print(f"  Trying fallback methods...")
        
        # Method 2: Try Aspose.Words (excellent quality, commercial)
        if HAS_ASPOSE:
            try:
                print(f"  Using Aspose.Words for conversion (excellent quality, pure Python)...")
                doc = aw.Document(temp_docx_path)
                doc.save(temp_pdf_path)
                
                # Read the converted PDF
                with open(temp_pdf_path, 'rb') as f:
                    pdf_content = f.read()
                return pdf_content
            except Exception as e:
                print(f"  Aspose.Words failed: {e}")
                print(f"  Trying next fallback method...")
        
        # Method 3: Try LibreOffice via command line (best free option)
        if HAS_SUBPROCESS:
            try:
                print(f"  Using LibreOffice for conversion...")
                
                # Try to find LibreOffice executable
                libreoffice_paths = [
                    'libreoffice',  # Linux standard
                    'soffice',  # Alternative name
                    '/usr/bin/libreoffice',  # Linux full path
                    '/usr/bin/soffice',  # Linux alternative full path
                    '/usr/lib/libreoffice/program/soffice',  # Some Linux distros
                    '/opt/libreoffice/program/soffice',  # Manual Linux installation
                    '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # Mac
                    'C:\\Program Files\\LibreOffice\\program\\soffice.exe',  # Windows
                ]
                
                libreoffice_cmd = None
                for path in libreoffice_paths:
                    try:
                        subprocess.run([path, '--version'], capture_output=True, timeout=5)
                        libreoffice_cmd = path
                        break
                    except:
                        continue
                
                if libreoffice_cmd:
                    # Convert using LibreOffice
                    cmd = [
                        libreoffice_cmd,
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', temp_dir,
                        temp_docx_path
                    ]
                    
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                    
                    if result.returncode == 0 and os.path.exists(temp_pdf_path):
                        # Read the converted PDF
                        with open(temp_pdf_path, 'rb') as f:
                            pdf_content = f.read()
                        return pdf_content
                    else:
                        print(f"  LibreOffice conversion failed with return code {result.returncode}")
                else:
                    print(f"  LibreOffice not found on system")
            except Exception as e:
                print(f"  LibreOffice conversion failed: {e}")
        
        # Method 4: Pure Python fallback with docx + fpdf2
        if HAS_DOCX_FPDF:
            try:
                print(f"  Using pure Python conversion (docx + fpdf2)...")
                print(f"  Note: This preserves text content but not complex formatting")
                
                # Read DOCX document
                doc = Document(temp_docx_path)
                
                # Create PDF
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Arial", size=11)
                
                # Add content from DOCX
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        # Check if it looks like a heading
                        if paragraph.style and paragraph.style.name and 'Heading' in paragraph.style.name:
                            pdf.set_font("Arial", 'B', 14)
                            pdf.multi_cell(0, 10, text)
                            pdf.set_font("Arial", size=11)
                        else:
                            # Regular paragraph
                            pdf.multi_cell(0, 6, text)
                        pdf.ln(2)
                
                # Add tables as text
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            pdf.multi_cell(0, 6, row_text)
                    pdf.ln(2)
                
                # Save PDF
                pdf.output(temp_pdf_path)
                
                # Read the converted PDF
                with open(temp_pdf_path, 'rb') as f:
                    pdf_content = f.read()
                return pdf_content
            except Exception as e:
                print(f"  Pure Python conversion failed: {e}")
        
        # Method 5: Try pypandoc as last resort (poor formatting preservation)
        if HAS_PYPANDOC:
            try:
                print(f"  WARNING: Using pypandoc (formatting may not be preserved)...")
                pypandoc.convert_file(
                    temp_docx_path,
                    'pdf',
                    outputfile=temp_pdf_path,
                    extra_args=['--pdf-engine=xelatex']  # Better Unicode support
                )
                
                # Read the converted PDF
                with open(temp_pdf_path, 'rb') as f:
                    pdf_content = f.read()
                print(f"  Note: pypandoc conversion completed but formatting may differ from original")
                return pdf_content
            except Exception as e:
                print(f"  pypandoc failed: {e}")
        
        # If all methods fail, inform user about installation options
        print(f"\n[ERROR] No suitable DOCX to PDF converter found!")
        print(f"Please install one of the following:")
        print(f"\nFor Dataiku/restricted environments (Python packages only):")
        print(f"  pip install python-docx fpdf2  # Basic but pure Python solution")
        print(f"  pip install aspose-words       # Commercial, excellent quality")
        print(f"\nFor full system access:")
        print(f"  1. LibreOffice from https://www.libreoffice.org/ (best free option)")
        print(f"  2. For Windows/Mac: pip install docx2pdf (requires MS Word)")
        return None
        
    except Exception as e:
        print(f"[ERROR] Failed to convert DOCX to PDF for '{filename}': {e}")
        return None
    finally:
        # Clean up temporary directory
        try:
            shutil.rmtree(temp_dir)
        except:
            pass

# ==============================================================================
# --- Main Execution Logic ---
# ==============================================================================

def main():
    print("\n" + "="*60)
    print("DOCX/PDF to PDF Preprocessor for Catalog Search Pipeline")
    print("="*60 + "\n")
    
    # Check available conversion methods
    print("[0] Checking available DOCX to PDF conversion methods...")
    system_platform = platform.system()
    print(f"  Platform detected: {system_platform}")
    
    converters_found = []
    
    # Check docx2pdf (not for Linux)
    if HAS_DOCX2PDF:
        if system_platform == "Linux":
            print("  ✗ docx2pdf installed but NOT usable on Linux (requires MS Word)")
            print("    Note: docx2pdf only works on Windows/Mac with Microsoft Word")
        else:
            converters_found.append("docx2pdf (excellent quality, MS Word required)")
            print("  ✓ docx2pdf available (excellent quality, preserves all formatting)")
    
    if HAS_ASPOSE:
        converters_found.append("Aspose.Words (excellent quality, pure Python)")
        print("  ✓ Aspose.Words available (excellent quality, commercial)")
    
    if HAS_SUBPROCESS:
        converters_found.append("LibreOffice (will check at runtime)")
        print("  ✓ LibreOffice command line available (checking at runtime)")
    
    if HAS_DOCX_FPDF:
        converters_found.append("docx + fpdf2 (pure Python, basic formatting)")
        print("  ✓ Pure Python conversion available (python-docx + fpdf2)")
        print("    Note: Preserves text content but not complex formatting")
    
    if HAS_PYPANDOC:
        converters_found.append("pypandoc (poor formatting preservation)")
        print("  ⚠ pypandoc Python package found but needs pandoc system package")
        print("    Install pandoc: sudo apt-get install pandoc (Ubuntu/Debian)")
        print("    Or: pypandoc.download_pandoc() in Python")
    
    if not converters_found or (system_platform == "Linux" and len(converters_found) == 1 and "LibreOffice" in converters_found[0]):
        print("\n" + "="*60)
        if system_platform == "Linux":
            print("IMPORTANT FOR LINUX USERS:")
            print("You need to install LibreOffice for DOCX conversion:")
            print("  sudo apt-get install libreoffice  (Ubuntu/Debian)")
            print("  sudo yum install libreoffice      (RHEL/CentOS)")
            print("  sudo dnf install libreoffice      (Fedora)")
            print("\nOptionally, if you want to use pypandoc:")
            print("  sudo apt-get install pandoc")
        else:
            print("Recommended installation (in order of preference):")
            print("  1. LibreOffice from https://www.libreoffice.org/ (best free option)")
            print("  2. For Windows/Mac: pip install docx2pdf (requires MS Word)")
            print("  3. For commercial use: pip install aspose-words (pure Python)")
        print("="*60)
    print("-" * 60)
    
    # Connect to NAS
    print("[1] Connecting to NAS...")
    conn = create_nas_connection()
    if not conn:
        print("[CRITICAL ERROR] Failed to connect to NAS. Exiting.")
        sys.exit(1)
    
    try:
        # List all DOCX and PDF files in input folder
        print(f"\n[2] Searching for DOCX and PDF files in: {NAS_PARAMS['share']}/{NAS_INPUT_FOLDER_PATH}")
        all_files = list_docx_pdf_files_from_nas(conn, NAS_PARAMS["share"], NAS_INPUT_FOLDER_PATH)
        
        if not all_files:
            print("[INFO] No DOCX or PDF files found in input folder.")
            print("="*60 + "\n")
            return
        
        # Separate PDFs and DOCX files
        pdf_files = [f for f in all_files if f['is_pdf']]
        docx_files = [f for f in all_files if not f['is_pdf']]
        
        print(f"Found {len(all_files)} files total:")
        print(f"  - PDF files: {len(pdf_files)}")
        print(f"  - DOCX/DOC files: {len(docx_files)}")
        
        # Process files
        print(f"\n[3] Processing files...")
        success_count = 0
        error_count = 0
        
        # Process PDF files (direct copy)
        if pdf_files:
            print(f"\n--- Processing {len(pdf_files)} PDF files ---")
            for file_info in pdf_files:
                source_path = file_info['path']
                filename = file_info['filename']
                dest_path = os.path.join(NAS_OUTPUT_FOLDER_PATH, filename).replace('\\', '/')
                
                print(f"\nCopying PDF: {filename}")
                print(f"  From: {source_path}")
                print(f"  To: {dest_path}")
                
                if copy_pdf_to_output(conn, NAS_PARAMS["share"], source_path, dest_path):
                    success_count += 1
                else:
                    error_count += 1
        
        # Process DOCX files (convert to PDF)
        if docx_files:
            print(f"\n--- Processing {len(docx_files)} DOCX/DOC files ---")
            for file_info in docx_files:
                source_path = file_info['path']
                filename = file_info['filename']
                
                # Change extension to .pdf for output
                pdf_filename = os.path.splitext(filename)[0] + '.pdf'
                dest_path = os.path.join(NAS_OUTPUT_FOLDER_PATH, pdf_filename).replace('\\', '/')
                
                print(f"\nConverting DOCX: {filename}")
                print(f"  Source: {source_path}")
                print(f"  Output: {dest_path}")
                
                try:
                    # Read DOCX content
                    docx_content = read_file_from_nas(conn, NAS_PARAMS["share"], source_path)
                    if docx_content is None:
                        print(f"[ERROR] Failed to read DOCX file: {filename}")
                        error_count += 1
                        continue
                    
                    # Convert to PDF
                    pdf_content = convert_docx_to_pdf(docx_content, filename)
                    if pdf_content is None:
                        print(f"[ERROR] Failed to convert DOCX to PDF: {filename}")
                        error_count += 1
                        continue
                    
                    # Save PDF to NAS
                    if save_pdf_to_nas(conn, NAS_PARAMS["share"], dest_path, pdf_content):
                        success_count += 1
                    else:
                        error_count += 1
                        
                except Exception as e:
                    print(f"[ERROR] Failed to process DOCX file '{filename}': {e}")
                    error_count += 1
        
        # Summary
        print("\n" + "="*60)
        print("Processing Complete!")
        print(f"Successfully processed: {success_count} files")
        if error_count > 0:
            print(f"Failed processing: {error_count} files")
        print("="*60 + "\n")
        
    finally:
        if conn:
            conn.close()
            print("NAS connection closed.")

if __name__ == "__main__":
    main()